import streamlit as st
from docx import Document
import io
import requests

# Function to upload file to Flask server
def upload_to_flask(file_name, file_data):
    url = 'http://<your-local-ip>:5000/upload'
    files = {'file': (file_name, file_data, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    response = requests.post(url, files=files)
    return response

# Create a Streamlit app
st.title("Word Saver App")

# Input text box
word = st.text_input("Enter a word")

# Button to save the word
if st.button("Save to Word Document"):
    if word:
        # Create a new Document
        doc = Document()
        doc.add_heading('Words List', level=1)
        doc.add_paragraph(word)

        # Save the document to a BytesIO object
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Upload the document to Flask server
        response = upload_to_flask("words.docx", buffer)
        if response.status_code == 200:
            st.success(f"The word '{word}' has been saved to the document on your PC.")
        else:
            st.error("Failed to save the document on your PC.")
    else:
        st.error("Please enter a word before saving.")
